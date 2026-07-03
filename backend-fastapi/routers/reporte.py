import os
from io import BytesIO
from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def set_cell_background(cell, fill_color):
    shd_xml = f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>'
    shd = parse_xml(shd_xml)
    cell._tc.get_or_add_tcPr().append(shd)

from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            if edge == 'start': tag = 'w:left'
            if edge == 'end': tag = 'w:right'
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def add_borders_to_table(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"sz": 4, "val": "single", "color": "000000"},
                bottom={"sz": 4, "val": "single", "color": "000000"},
                start={"sz": 4, "val": "single", "color": "000000"},
                end={"sz": 4, "val": "single", "color": "000000"}
            )




from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch

import models
from database import get_db

router = APIRouter()

def draw_background(canvas, doc):
    canvas.saveState()
    width, height = letter
    
    # Draw blue decorative shapes in top right
    path = canvas.beginPath()
    path.moveTo(width, height)
    path.lineTo(width, height - 150)
    path.lineTo(width - 150, height)
    canvas.setFillColor(colors.HexColor('#dbe4f0')) # light blue
    canvas.setStrokeColor(colors.HexColor('#dbe4f0'))
    canvas.drawPath(path, fill=1, stroke=0)

    path2 = canvas.beginPath()
    path2.moveTo(width, height - 50)
    path2.lineTo(width, height - 180)
    path2.lineTo(width - 50, height - 180)
    canvas.setFillColor(colors.HexColor('#c3d1e6')) # slightly darker blue
    canvas.setStrokeColor(colors.HexColor('#c3d1e6'))
    canvas.drawPath(path2, fill=1, stroke=0)
    
    # Text "MICHAEL CESAR QUISPE SEBASTIAN" top right
    canvas.setFont("Helvetica-Bold", 10)
    canvas.setFillColor(colors.HexColor('#748bb6'))
    canvas.drawRightString(width - 30, height - 40, "MICHAEL CESAR QUISPE SEBASTIAN")
    
    # Logo top left
    logo_path = os.path.join(os.path.dirname(__file__), '..', 'static', 'assets', 'logo-mqs.png')
    if os.path.exists(logo_path):
        # Increased logo size and positioned it nicely
        canvas.drawImage(logo_path, 30, height - 90, width=2.5*inch, height=1*inch, preserveAspectRatio=True, mask='auto')
        
    # Title "LIQUIDACIÓN DE PRIMAS" centered with shadow
    canvas.setFont("Helvetica-Bold", 26)
    text = "LIQUIDACIÓN DE PRIMAS"
    text_width = canvas.stringWidth(text, "Helvetica-Bold", 26)
    x = (width - text_width) / 2
    y = height - 140
    
    # Shadow
    canvas.setFillColor(colors.HexColor('#9ca3af'))
    canvas.drawString(x + 2, y - 2, text)
    # Text
    canvas.setFillColor(colors.black)
    canvas.drawString(x, y, text)
    
    # --- FOOTER ---
    # Horizontal line
    canvas.setStrokeColor(colors.HexColor('#a7c0de'))
    canvas.setLineWidth(1)
    canvas.line(40, 75, width - 40, 75)
    
    # Decorative little flourish could be drawn or placed via image, but we'll approximate 
    # drawing a small asterisk or keeping it simple text for now to avoid needing an image asset.
    canvas.setFont("Helvetica", 14)
    canvas.setFillColor(colors.HexColor('#c3d1e6'))
    canvas.drawCentredString(width / 2, 63, "⚘") # Small decorative icon
    
    # Footer Text
    canvas.setFont("Helvetica-Bold", 8)
    canvas.setFillColor(colors.gray)
    canvas.drawCentredString(width/2, 50, "JR. CHIRA MZ. U LOTE 01 P.J. SAN LUIS AMARILIS - HUANUCO - HUANUCO")
    
    # "CORREO: MICHAELCQS@HOTMAIL.COM" with mixed colors
    text_correo1 = "CORREO: "
    text_correo2 = "MICHAELCQS@HOTMAIL.COM"
    w1 = canvas.stringWidth(text_correo1, "Helvetica-Bold", 8)
    w2 = canvas.stringWidth(text_correo2, "Helvetica-Bold", 8)
    total_w = w1 + w2
    x_start = (width - total_w) / 2
    canvas.drawString(x_start, 40, text_correo1)
    
    canvas.setFillColor(colors.HexColor('#748bb6'))
    canvas.drawString(x_start + w1, 40, text_correo2)
    # underline
    canvas.setStrokeColor(colors.HexColor('#748bb6'))
    canvas.setLineWidth(0.5)
    canvas.line(x_start + w1, 39, x_start + w1 + w2, 39)
    
    canvas.setFillColor(colors.gray)
    canvas.drawCentredString(width/2, 30, "CELULAR:932030030 - RUC 10423117864")
    
    canvas.restoreState()

@router.get("/empresa/{empresa_id}")
def descargar_reporte_pdf(empresa_id: int, db: Session = Depends(get_db)):
    # Fetch data
    empresa_model = db.query(models.Empresa).filter(models.Empresa.id == empresa_id).first()
    if not empresa_model:
        raise HTTPException(status_code=404, detail="Empresa no encontrada")
    
    fianzas = [
        {
            'numero': f.numero,
            'tipo': f.tipo,
            'monto': float(f.monto),
            'fecha_inicio': f.fecha_inicio,
            'fecha_vencimiento': f.fecha_vencimiento,
            'observaciones': f.observaciones
        }
        for f in db.query(models.CartaFianza).filter(models.CartaFianza.empresa_id == empresa_id).all()
    ]
    
    facturas = [
        {
            'id': f.id,
            'numero': f.numero,
            'monto': float(f.monto),
            'fecha_salida': f.fecha_salida,
            'tipo_fianza_relacionada': f.tipo_fianza_relacionada,
            'numero_fianza_relacionada': f.numero_fianza_relacionada,
            'observacion': f.observacion
        }
        for f in db.query(models.Factura).filter(models.Factura.empresa_id == empresa_id).all()
    ]
    
    consorciados = db.query(models.Consorciado).filter(models.Consorciado.empresa_id == empresa_id).all()
    consorciados_nombres = [c.nombre for c in consorciados]
    
    if len(consorciados_nombres) == 0:
        texto_consorcio = empresa_model.nombre
    elif len(consorciados_nombres) == 1:
        texto_consorcio = f"{empresa_model.nombre} – {consorciados_nombres[0]}"
    else:
        consorciados_str = ", ".join(consorciados_nombres[:-1]) + " y " + consorciados_nombres[-1]
        texto_consorcio = f"{empresa_model.nombre} – {consorciados_str}"
        
    buffer = BytesIO()
    # Portrait layout, start content lower to leave room for the custom header
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=40, leftMargin=40, topMargin=170, bottomMargin=90)
    styles = getSampleStyleSheet()
    
    # Custom Styles
    normal_style = styles['Normal']
    normal_style.fontSize = 10
    normal_style.leading = 14
    
    bold_style = ParagraphStyle('BoldStyle', parent=normal_style, fontName='Helvetica-Bold')
    
    elements = []
    
    # --- Introductory Paragraphs ---
    # Notice the bold is strictly placed just like the user requested.
    p1 = Paragraph(f"Por medio de la presente, se adjunta la <b>liquidación de primas emitidas por CESCE PERÚ S.A. COMPAÑÍA DE SEGUROS</b> a favor del {texto_consorcio}.", normal_style)
    elements.append(p1)
    elements.append(Spacer(1, 0.15 * inch))
    
    p2 = Paragraph("La documentación comprende desde la <b>Carta Fianza inicial por Fiel Cumplimiento y Adelanto Directo</b>, hasta la última renovación, conforme al cuadro y detalle adjunto.", normal_style)
    elements.append(p2)
    elements.append(Spacer(1, 0.2 * inch))

    # --- Resumen Financiero ---
    monto_obra = float(empresa_model.monto_obra) if empresa_model.monto_obra else 0.0
    suma_asegurada = float(empresa_model.suma_asegurada) if empresa_model.suma_asegurada else (monto_obra * 0.10)
    monto_garantia = float(empresa_model.monto_garantia) if empresa_model.monto_garantia else (suma_asegurada * 0.20)
    
    resumen_data = [
        ["RESUMEN FINANCIERO", "", ""],
        ["MONTO ADJUDICADO (S/)", "SUMA ASEGURADA (10%)", "FONDO GARANTÍA (20%)"],
        [f"S/ {monto_obra:,.2f}", f"S/ {suma_asegurada:,.2f}", f"S/ {monto_garantia:,.2f}"]
    ]
    
    t_resumen = Table(resumen_data, colWidths=[2*inch, 2*inch, 2*inch])
    t_resumen.setStyle(TableStyle([
        ('SPAN', (0, 0), (-1, 0)),
        ('BACKGROUND', (0, 0), (-1, 1), colors.HexColor("#1a365d")),
        ('TEXTCOLOR', (0, 0), (-1, 1), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 1), 6),
        ('BACKGROUND', (0, 2), (-1, 2), colors.HexColor("#f8fafc")),
        ('TEXTCOLOR', (0, 2), (-1, 2), colors.HexColor("#1e3a8a")),
        ('FONTNAME', (0, 2), (-1, 2), 'Helvetica-Bold'),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor("#cbd5e1")),
    ]))
    elements.append(t_resumen)
    elements.append(Spacer(1, 0.3 * inch))
    
    # --- Tables ---
    def get_invoice_type(f):
        if f.get('tipo_fianza_relacionada'): return f.get('tipo_fianza_relacionada')
        if not f.get('numero_fianza_relacionada'): return 'Fiel Cumplimiento'
        match = next((b for b in fianzas if b['numero'].strip() == f.get('numero_fianza_relacionada', '').strip()), None)
        return match['tipo'] if match else 'Fiel Cumplimiento'

    tipo_map = {
        'Fiel Cumplimiento': 'FC',
        'Adelanto Directo': 'AD',
        'Adelanto de Materiales': 'AM'
    }
    
    fianza_types = ['Fiel Cumplimiento', 'Adelanto Directo', 'Adelanto de Materiales']
    
    for f_type in fianza_types:
        fianzas_of_type = [f for f in fianzas if f['tipo'] == f_type]
        facturas_of_type = [f for f in facturas if get_invoice_type(f) == f_type]
        
        if fianzas_of_type or facturas_of_type:
            poliza_abrev = tipo_map.get(f_type, 'FC')
            
            # Group facturas by fianza
            table_data = []
            
            # Title Row (spans 5 cols)
            table_data.append([empresa_model.nombre.upper(), "", "", "", ""])
            # Headers
            table_data.append(["N° FIANZA", "POLIZA", "MONTO", "N° FACTURA", "MONTO"])
            
            processed_facturas = set()
            total_facturas = 0.0
            
            # Add fianzas and their facturas
            for fianza in fianzas_of_type:
                rel_facturas = [f for f in facturas_of_type if f.get('numero_fianza_relacionada') and f.get('numero_fianza_relacionada').strip() == fianza['numero'].strip()]
                
                if not rel_facturas:
                    table_data.append([
                        fianza['numero'],
                        poliza_abrev,
                        f"S/ {fianza['monto']:,.2f}",
                        "",
                        ""
                    ])
                else:
                    for fact in rel_facturas:
                        processed_facturas.add(fact['id'])
                        total_facturas += fact['monto']
                        table_data.append([
                            fianza['numero'],
                            poliza_abrev,
                            f"S/ {fianza['monto']:,.2f}",
                            fact['numero'],
                            f"S/ {fact['monto']:,.2f}"
                        ])
            
            # Orphaned facturas
            for fact in facturas_of_type:
                if fact['id'] not in processed_facturas:
                    total_facturas += fact['monto']
                    table_data.append([
                        "---",
                        poliza_abrev,
                        "---",
                        fact['numero'],
                        f"S/ {fact['monto']:,.2f}"
                    ])
            
            # Total Row
            table_data.append(["TOTAL", "", "", "", f"S/ {total_facturas:,.2f}"])
            
            col_widths = [1.7*inch, 0.7*inch, 1.2*inch, 1.2*inch, 1.1*inch]
            
            t = Table(table_data, colWidths=col_widths)
            
            t_style = TableStyle([
                # Title Row
                ('SPAN', (0, 0), (-1, 0)),
                ('BACKGROUND', (0, 0), (-1, 1), colors.HexColor("#dbe4f0")),
                ('FONTNAME', (0, 0), (-1, 1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 1), 9),
                ('ALIGN', (0, 0), (-1, 1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, 1), 'MIDDLE'),
                
                # Data Rows
                ('FONTNAME', (0, 2), (-1, -2), 'Helvetica'),
                ('FONTSIZE', (0, 2), (-1, -2), 9),
                ('ALIGN', (0, 2), (-1, -2), 'CENTER'),
                ('VALIGN', (0, 2), (-1, -2), 'MIDDLE'),
                
                # Total Row
                ('SPAN', (0, -1), (3, -1)),
                ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, -1), (-1, -1), 9),
                ('ALIGN', (0, -1), (3, -1), 'CENTER'),
                ('ALIGN', (4, -1), (4, -1), 'CENTER'),
                
                # Grid
                ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                ('PADDING', (0, 0), (-1, -1), 4),
            ])
            t.setStyle(t_style)
            
            elements.append(t)
            elements.append(Spacer(1, 0.4 * inch))

    # Build PDF with custom first page background
    doc.build(elements, onFirstPage=draw_background, onLaterPages=draw_background)
    buffer.seek(0)
    
    headers = {
        'Content-Disposition': f'attachment; filename="Liquidacion_Primas_{empresa_model.nombre.replace(" ", "_")}.pdf"'
    }
    return StreamingResponse(buffer, headers=headers, media_type='application/pdf')


@router.get("/empresa/{empresa_id}/word")
def descargar_reporte_word(empresa_id: int, db: Session = Depends(get_db)):
    empresa_model = db.query(models.Empresa).filter(models.Empresa.id == empresa_id).first()
    if not empresa_model:
        raise HTTPException(status_code=404, detail="Empresa no encontrada")
        
    fianzas = [
        {
            'numero': f.numero,
            'tipo': f.tipo,
            'monto': float(f.monto),
            'fecha_inicio': f.fecha_inicio,
            'fecha_vencimiento': f.fecha_vencimiento,
            'observaciones': f.observaciones
        }
        for f in db.query(models.CartaFianza).filter(models.CartaFianza.empresa_id == empresa_id).all()
    ]
    
    facturas = [
        {
            'id': f.id,
            'numero': f.numero,
            'monto': float(f.monto),
            'fecha_salida': f.fecha_salida,
            'tipo_fianza_relacionada': f.tipo_fianza_relacionada,
            'numero_fianza_relacionada': f.numero_fianza_relacionada,
            'observacion': f.observacion
        }
        for f in db.query(models.Factura).filter(models.Factura.empresa_id == empresa_id).all()
    ]
    
    consorciados = db.query(models.Consorciado).filter(models.Consorciado.empresa_id == empresa_id).all()
    consorciados_nombres = [c.nombre for c in consorciados]
    
    if len(consorciados_nombres) == 0:
        texto_consorcio = empresa_model.nombre
    elif len(consorciados_nombres) == 1:
        texto_consorcio = f"{empresa_model.nombre} – {consorciados_nombres[0]}"
    else:
        consorciados_str = ", ".join(consorciados_nombres[:-1]) + " y " + consorciados_nombres[-1]
        texto_consorcio = f"{empresa_model.nombre} – {consorciados_str}"

    template_path = os.path.join(os.path.dirname(__file__), '..', '..', 'plantilla_primas.docx')
    if os.path.exists(template_path):
        document = DocxDocument(template_path)
        # Limpiar cualquier párrafo vacío o contenido basura en el cuerpo de la plantilla 
        # para que nuestro contenido empiece en la página 1.
        for p in list(document.paragraphs):
            p._element.getparent().remove(p._element)
    else:
        document = DocxDocument()
        
    if not os.path.exists(template_path):
        heading = document.add_paragraph("LIQUIDACIÓN DE PRIMAS")
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.name = 'Helvetica'
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.bold = True
        document.add_paragraph()
    
    # Intro
    p = document.add_paragraph()
    p.add_run("Por medio de la presente, se adjunta la ")
    r2 = p.add_run("liquidación de primas emitidas por CESCE PERÚ S.A. COMPAÑÍA DE SEGUROS")
    r2.bold = True
    p.add_run(f" a favor del {texto_consorcio}.\n\n")
    p.add_run("La documentación comprende desde la ")
    r3 = p.add_run("Carta Fianza inicial por Fiel Cumplimiento y Adelanto Directo")
    r3.bold = True
    p.add_run(", hasta la última renovación, conforme al cuadro y detalle adjunto.")
    
    # Resumen Financiero
    monto_obra = float(empresa_model.monto_obra) if empresa_model.monto_obra else 0.0
    suma_asegurada = float(empresa_model.suma_asegurada) if empresa_model.suma_asegurada else (monto_obra * 0.10)
    monto_garantia = float(empresa_model.monto_garantia) if empresa_model.monto_garantia else (suma_asegurada * 0.20)
    
    document.add_paragraph()
    
    table = document.add_table(rows=3, cols=3)
    try:
        table.style = 'Table Grid'
    except KeyError:
        pass
    
    cell0 = table.cell(0, 0)
    cell0.merge(table.cell(0, 2))
    cell0.text = "RESUMEN FINANCIERO"
    set_cell_background(cell0, "1a365d")
    if cell0.paragraphs[0].runs:
        cell0.paragraphs[0].runs[0].bold = True
        cell0.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    table.cell(1, 0).text = "MONTO ADJUDICADO (S/)"
    table.cell(1, 1).text = "SUMA ASEGURADA (10%)"
    table.cell(1, 2).text = "FONDO GARANTÍA (20%)"
    
    set_cell_background(table.cell(1, 0), "1a365d")
    set_cell_background(table.cell(1, 1), "1a365d")
    set_cell_background(table.cell(1, 2), "1a365d")
    
    for i in range(3):
        if table.cell(1, i).paragraphs[0].runs:
            table.cell(1, i).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    table.cell(2, 0).text = f"S/ {monto_obra:,.2f}"
    table.cell(2, 1).text = f"S/ {suma_asegurada:,.2f}"
    table.cell(2, 2).text = f"S/ {monto_garantia:,.2f}"
    
    for row in table.rows:
        for cell in row.cells:
            for p_elem in cell.paragraphs:
                p_elem.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p_elem.runs:
                    run.font.size = Pt(9)
                    if row._index < 2:
                        run.bold = True
    
    add_borders_to_table(table)
    document.add_paragraph()
    
    # Fianzas vs Facturas
    def get_invoice_type(f):
        if f.get('tipo_fianza_relacionada'): return f.get('tipo_fianza_relacionada')
        if not f.get('numero_fianza_relacionada'): return 'Fiel Cumplimiento'
        match = next((b for b in fianzas if b['numero'].strip() == f.get('numero_fianza_relacionada', '').strip()), None)
        return match['tipo'] if match else 'Fiel Cumplimiento'

    tipo_map = {
        'Fiel Cumplimiento': 'FC',
        'Adelanto Directo': 'AD',
        'Adelanto de Materiales': 'AM'
    }
    
    fianza_types = ['Fiel Cumplimiento', 'Adelanto Directo', 'Adelanto de Materiales']
    
    for f_type in fianza_types:
        fianzas_of_type = [f for f in fianzas if f['tipo'] == f_type]
        facturas_of_type = [f for f in facturas if get_invoice_type(f) == f_type]
        
        if fianzas_of_type or facturas_of_type:
            poliza_abrev = tipo_map.get(f_type, 'FC')
            
            t = document.add_table(rows=2, cols=5)
            try:
                t.style = 'Table Grid'
            except KeyError:
                pass
            
            c0 = t.cell(0, 0)
            c0.merge(t.cell(0, 4))
            c0.text = empresa_model.nombre.upper()
            set_cell_background(c0, "dbe4f0")
            if c0.paragraphs[0].runs:
                c0.paragraphs[0].runs[0].bold = True
            
            headers = ["N° FIANZA", "POLIZA", "MONTO", "N° FACTURA", "MONTO"]
            for i, h in enumerate(headers):
                t.cell(1, i).text = h
                if t.cell(1, i).paragraphs[0].runs:
                    t.cell(1, i).paragraphs[0].runs[0].bold = True
            
            processed_facturas = set()
            total_facturas = 0.0
            
            for fianza in fianzas_of_type:
                rel_facturas = [f for f in facturas_of_type if f.get('numero_fianza_relacionada') and f.get('numero_fianza_relacionada').strip() == fianza['numero'].strip()]
                
                if not rel_facturas:
                    row = t.add_row()
                    row.cells[0].text = fianza['numero']
                    row.cells[1].text = poliza_abrev
                    row.cells[2].text = f"S/ {fianza['monto']:,.2f}"
                else:
                    for fact in rel_facturas:
                        processed_facturas.add(fact['id'])
                        total_facturas += fact['monto']
                        row = t.add_row()
                        row.cells[0].text = fianza['numero']
                        row.cells[1].text = poliza_abrev
                        row.cells[2].text = f"S/ {fianza['monto']:,.2f}"
                        row.cells[3].text = fact['numero']
                        row.cells[4].text = f"S/ {fact['monto']:,.2f}"
            
            for fact in facturas_of_type:
                if fact['id'] not in processed_facturas:
                    total_facturas += fact['monto']
                    row = t.add_row()
                    row.cells[0].text = "---"
                    row.cells[1].text = poliza_abrev
                    row.cells[2].text = "---"
                    row.cells[3].text = fact['numero']
                    row.cells[4].text = f"S/ {fact['monto']:,.2f}"
            
            row = t.add_row()
            cell_total = row.cells[0]
            cell_total.merge(row.cells[3])
            cell_total.text = "TOTAL"
            if cell_total.paragraphs[0].runs:
                cell_total.paragraphs[0].runs[0].bold = True
            
            cell_val = row.cells[4]
            cell_val.text = f"S/ {total_facturas:,.2f}"
            if cell_val.paragraphs[0].runs:
                cell_val.paragraphs[0].runs[0].bold = True
            
            for row in t.rows:
                for cell in row.cells:
                    for p_elem in cell.paragraphs:
                        p_elem.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p_elem.runs:
                            run.font.size = Pt(9)
            
            add_borders_to_table(t)
            document.add_paragraph()

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    headers_dict = {
        'Content-Disposition': f'attachment; filename="Liquidacion_Primas_{empresa_model.nombre.replace(" ", "_")}.docx"'
    }
    
    return StreamingResponse(buffer, headers=headers_dict, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
